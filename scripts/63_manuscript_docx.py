# -*- coding: utf-8 -*-
"""63_manuscript_docx.py - 把主稿 markdown 转为格式化 docx"""
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1] if len(sys.argv) > 1 else "docs/manuscript_main_20260901.md"
OUT = sys.argv[2] if len(sys.argv) > 2 else "docs/manuscript_main_20260901.docx"

doc = Document()

# 全局样式
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(11)
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(6)

for i in (1, 2, 3):
    hs = doc.styles["Heading %d" % i]
    hs.font.name = "Times New Roman"
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
doc.styles["Heading 1"].font.size = Pt(14)
doc.styles["Heading 2"].font.size = Pt(12)
doc.styles["Heading 3"].font.size = Pt(11)


def add_runs(p, text):
    """解析 **bold** 与 *italic* 行内标记"""
    # 先按 ** 分段
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1])
            r.italic = True
        else:
            p.add_run(part)


def is_table_line(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def is_sep_line(line):
    return bool(re.match(r"^\|[\s:|-]+\|$", line.strip()))


def add_table(rows):
    """rows: list of list of cell strings (含表头行)"""
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            txt = row[j] if j < len(row) else ""
            r = p.add_run(txt)
            r.font.size = Pt(9)
            r.font.name = "Times New Roman"
            if i == 0:
                r.bold = True


with open(SRC, encoding="utf-8") as f:
    lines = f.read().split("\n")

# 预处理: 把表格行块合并为单个表格元素
elements = []  # ("para", line) 或 ("table", rows)
buf_rows = []
for line in lines:
    if is_sep_line(line):
        # 分隔行: 跳过, 不打断当前表格块
        continue
    if is_table_line(line):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        buf_rows.append(cells)
    else:
        if buf_rows:
            elements.append(("table", buf_rows))
            buf_rows = []
        elements.append(("para", line))
if buf_rows:
    elements.append(("table", buf_rows))

first_h1_done = False
for kind, payload in elements:
    if kind == "table":
        add_table(payload)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        continue
    line = payload.rstrip()
    if not line.strip():
        continue
    if line.strip() == "---":
        continue
    if line.startswith("# ") and not first_h1_done:
        # 标题: 居中, 14pt
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line[2:].strip())
        r.bold = True
        r.font.size = Pt(14)
        first_h1_done = True
        continue
    if line.startswith("### "):
        doc.add_heading("", level=3)
        p = doc.paragraphs[-1]
        add_runs(p, line[4:].strip())
        continue
    if line.startswith("## "):
        doc.add_heading("", level=2)
        p = doc.paragraphs[-1]
        add_runs(p, line[3:].strip())
        continue
    if line.startswith("# "):
        doc.add_heading("", level=1)
        p = doc.paragraphs[-1]
        add_runs(p, line[2:].strip())
        continue
    # 作者/单位/通讯行: 在 Abstract 之前的非标题行居中
    p = doc.add_paragraph()
    if not first_h1_done or re.match(r"^\d+ |^\*Correspondence", line):
        pass
    add_runs(p, line.strip())

doc.save(OUT)
print("saved", OUT)
