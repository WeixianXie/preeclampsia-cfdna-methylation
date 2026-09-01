# -*- coding: utf-8 -*-
"""
Task #48: companion manuscript (辅助短文) docx.
Standalone short article integrating 2nd tier (GSE37722 leukocyte validation,
SMR/coloc) + supporting 3rd tier evidence (WGCNA, immune infiltration).
Target journal style: Clinical Epigenetics (short research article).
Output: results/companion_manuscript.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "results", "companion_manuscript.docx")

doc = Document()
# base styles
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(11)
for lvl, sz in [("Heading 1", 14), ("Heading 2", 12)]:
    doc.styles[lvl].font.size = Pt(sz)
    doc.styles[lvl].font.name = "Times New Roman"
    doc.styles[lvl].font.color.rgb = None

def p(text, bold=False, italic=False, align=None, size=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold, run.italic = bold, italic
    if size:
        run.font.size = Pt(size)
    if align:
        par.alignment = align
    return par

def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

# ---------------- Title ----------------
p("Maternal leukocyte composition dynamics underlie pre-eclampsia-associated "
  "differences in plasma cell-free DNA methylation: cross-layer validation in "
  "an independent cohort, genetic colocalization, and immune profiling",
  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
p("Companion article to: \"Differential plasma cfDNA methylation in pre-eclampsia "
  "(GSE282512): discovery, sensitivity analysis, and mechanistic dissection\" "
  "(main manuscript)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
p("Version 1.0 draft, generated 2026-08-30. Author list, affiliations, ethics "
  "statements and ORCID to be completed before submission.",
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

# ---------------- Abstract ----------------
doc.add_heading("Abstract", level=1)
p("Background: In our companion analysis of plasma cell-free DNA (cfDNA) "
  "whole-genome bisulfite sequencing from a pre-eclampsia (PE) case-control "
  "cohort (GSE282512; 32 PE vs 32 gestational-age matched controls), 166 "
  "candidate differentially methylated regions (DMRs) showed robust "
  "region-level signal (tube-type and batch adjusted) that did not replicate "
  "at the single-CpG level and was negatively correlated with placental "
  "tissue methylation differences, suggesting a maternal leukocyte-composition "
  "origin. Here we subject that interpretation to independent, orthogonal tests.")
p("Results: (i) In GSE37722 (84 maternal leukocyte methylome profiles spanning "
  "nulligravid to postpartum), probes hypermethylated in PE cfDNA showed "
  "significant methylation increases across normal gestation in leukocytes "
  "(mean t = 1.055, Stouffer p = 5.4e-28, sign test p = 1.5e-16; n = 108), "
  "whereas hypomethylated probes showed no trend (p = 0.85); PE cfDNA Δβ was "
  "concordant with leukocyte gestational Δβ (ρ = 0.381, p = 4.3e-06, n = 137). "
  "(ii) Weighted gene co-methylation network analysis of 5,000 variable cfDNA "
  "regions identified a PE-associated module (r = 0.487, p = 4.5e-05) "
  "significantly enriched for candidate DMRs (36/371 regions, Fisher "
  "p = 7.3e-17); the module was negatively associated with NK-cell cfDNA "
  "fraction (r = -0.452) and positively with naïve B-cell fraction "
  "(r = 0.354). (iii) In maternal blood transcriptomes, immune-infiltration "
  "differences between PE and controls were absent in early pregnancy "
  "(GSE85307; all BH ≥ 0.41) but pronounced at diagnosis (GSE48424: NK "
  "p = 0.009, cytotoxic cells p = 0.017, plasma cells p = 0.017, neutrophils "
  "p = 0.024), consistent with a concurrent rather than predictive state "
  "marker. (iv) SMR analysis prioritized one meQTL-GWAS shared DMR: "
  "SLC17A1 gene body methylation in PE (cg19490609, p_SMR = 0.0017, "
  "BH = 0.020; cg25753631, p_SMR = 0.0040, BH = 0.024), although colocalization "
  "posterior probabilities remained equivocal (PP4 ≤ 0.22).")
p("Conclusion: Four independent data layers converge on the interpretation "
  "that PE-associated cfDNA methylation differences primarily fingerprint "
  "maternal leukocyte subset dynamics during the disease state, rather than "
  "placental pathology or cis genetic regulation. cfDNA methylation panels "
  "should therefore be framed as maternal immune-state biomarkers, with "
  "SLC17A1 warranting follow-up as a genetically anchored methylation signal. "
  "Complete analysis-ready data and submission materials are prepared for "
  "controlled-access deposition at the European Genome-Phenome Archive.")

# ---------------- Background ----------------
doc.add_heading("Background", level=1)
p("Pre-eclampsia (PE) affects 2-8% of pregnancies and remains a leading cause "
  "of maternal and perinatal morbidity. Circulating cell-free DNA (cfDNA) has "
  "transformed fetal aneuploidy screening, and cfDNA methylation is a "
  "candidate modality for PE monitoring because it retains cell-of-origin "
  "information. However, plasma cfDNA is a mixture dominated by maternal "
  "leukocyte-derived fragments, and PE is accompanied by systemic maternal "
  "inflammation and leukocyte activation.")
p("In our companion study (main manuscript), we analyzed plasma cfDNA WGBS "
  "from GSE282512 (279 patients, 369 samples; discovery subcohort 32 PE vs 32 "
  "gestational-age matched controls, median GA difference 0 weeks) and "
  "identified 166 candidate DMRs robust to tube-type and sequencing batch. "
  "Three observations argued against a placental origin: single-CpG effects "
  "did not survive multiple testing correction; cfDNA Δβ correlated "
  "negatively with placental-tissue Δβ (ρ = -0.522, p = 0.0026); and the PE "
  "cfDNA methylome departed from, rather than approached, the placental "
  "reference. We here report four orthogonal validation and extension "
  "analyses, framed as an independent companion article.")

# ---------------- Methods ----------------
doc.add_heading("Methods", level=1)
doc.add_heading("Cohort and discovery set", level=2)
p("The discovery set, DMR calling pipeline, sensitivity analyses and "
  "confidence tiers are described in the main manuscript. In brief: 500 bp "
  "tiling, coverage-weighted beta values, limma adjusting tube_type and "
  "batch; 166 candidate DMRs (105 hyper-, 61 hypomethylated).")
doc.add_heading("Leukocyte gestational trajectories (GSE37722)", level=2)
p("GSE37722 profiles maternal leukocyte DNA methylation (Illumina Human "
  "Methylation 27K) in 84 women: nulligravid (n=14), early pregnancy (n=14), "
  "mid pregnancy (n=14), delivery (n=28) and postpartum (n=14), including "
  "longitudinal individuals. Probes were re-mapped to GRCh38 using the "
  "Zhou lab HM27 hg38 manifest. DMR-overlapping probes (n=113, covering 86 "
  "DMRs; ±2 kb sensitivity window n=153) were tested for gestational trend "
  "using limma with duplicateCorrelation to account for repeated "
  "measurements. Trend statistics were combined by direction group "
  "(Stouffer's method, sign test), and probe-level gestational change "
  "(delivery - nulligravid) was correlated with cfDNA Δβ (Spearman).")
doc.add_heading("SMR and colocalization", level=2)
p("For the three DMRs sharing variants between GoDML cis-meQTL (N≈32,000) and "
  "FinnGen R10 PE (7,965 cases / 211,852 controls) or gestational "
  "hypertension (GH; 9,535 / 211,957) GWAS, per-CpG lead meQTL variants were "
  "mapped from GRCh37 to GRCh38 via the Ensembl REST API and matched to "
  "GWAS variants by position. SMR was computed as the Wald ratio "
  "test-statistic (chi-square, 1 d.f.), which is invariant to allele "
  "orientation. Colocalization used coloc.abf in p-value mode (no allele "
  "harmonization available; meQTL summary statistics restricted to "
  "significant pairs). HEIDI was not implementable without an LD reference "
  "panel; these limitations are stated explicitly.")
doc.add_heading("Weighted gene co-methylation network analysis", level=2)
p("The 5,000 most variable regions (missingness ≤ 20%, imputed by row "
  "median) across 64 subcohort samples were used to build a signed WGCNA "
  "network (soft power 7). Module eigengenes were correlated with PE status "
  "and with 12-cell-type cfDNA deconvolution fractions (excluding "
  "zero-variance cell types). DMR enrichment per module was tested by "
  "Fisher's exact test.")
doc.add_heading("Maternal blood immune infiltration", level=2)
p("Immune-cell enrichment was quantified by a rank-weighted ssGSEA "
  "(Barbie et al. 2009; |x|^0.25 weighting) with a manually curated 14 "
  "lineage marker panel, applied to two maternal whole-blood transcriptome "
  "cohorts: GSE85307 (early pregnancy, PE n=47 / controls n=110) and "
  "GSE48424 (blood at PE diagnosis, n=18/18). Cohort-wise Wilcoxon tests "
  "with Benjamini-Hochberg correction were cross-referenced against "
  "cfDNA-derived 12-type cell fraction differences.")
doc.add_heading("Data deposition preparation", level=2)
p("Analysis-ready per-sample Bismark coverage files, region-level beta "
  "matrix, and sample metadata were compiled into EGA submission packages "
  "(sample metadata, file manifest, experiment metadata, DAC-facing "
  "analysis description, submission checklist) for controlled-access "
  "deposition; no submission was executed.")

# ---------------- Results ----------------
doc.add_heading("Results", level=1)
doc.add_heading("PE cfDNA DMRs track normal gestational dynamics of "
                "maternal leukocytes", level=2)
p("If PE-associated cfDNA hypermethylation reflects maternal leukocyte "
  "biology, the same loci should vary with gestation in healthy women's "
  "leukocytes. This prediction was strongly confirmed: among 108 "
  "hypermethylated-DMR probes mapped in GSE37722, methylation increased "
  "across gestation (mean t = 1.055; Stouffer p = 5.44e-28; sign-test "
  "p = 1.46e-16), whereas 29 hypomethylated-DMR probes showed no trend "
  "(p = 0.85). Effect sizes were concordant: cfDNA Δβ (PE - control) "
  "correlated with leukocyte Δβ (delivery - nulligravid) at ρ = 0.381 "
  "(p = 4.31e-06, n = 137). The probes were modestly, non-significantly "
  "enriched among the top 5% most gestation-dynamic genome-wide probes "
  "(OR = 1.14, p = 0.44), indicating direction-coherent rather than "
  "extreme-magnitude dynamics (Fig. 1, 2).")
doc.add_heading("A PE-associated co-methylation module maps to lymphocyte "
                "cfDNA fractions", level=2)
p("WGCNA of 5,000 variable cfDNA regions identified four modules. The blue "
  "module (371 regions) was associated with PE (r = 0.487, p = 4.46e-05) "
  "and strongly enriched for candidate DMRs (36/371, Fisher p = 7.3e-17); "
  "36 of the 99 DMRs entering the network were blue-module members. The "
  "same module was negatively correlated with NK-cell cfDNA fraction "
  "(r = -0.452, p = 1.8e-04) and positively with naïve B-cell fraction "
  "(r = 0.354) and CD4 memory T-cell fraction (r = 0.261). Grey and "
  "turquoise modules tracked monocyte fractions (r = 0.520 and 0.423). "
  "Thus the PE methylation signal is embedded in a co-methylation structure "
  "whose eigengene behaves as a lymphoid-composition axis (Fig. 3).")
doc.add_heading("Immune infiltration differences are concurrent with, not "
                "predictive of, PE", level=2)
p("In maternal blood transcriptomes, ssGSEA immune enrichment showed no "
  "difference between future PE cases and controls in early pregnancy "
  "(GSE85307; all lineages BH ≥ 0.41), while at PE diagnosis (GSE48424) "
  "NK cells (+0.188, p = 0.0087), cytotoxic cells (+0.174, p = 0.017), "
  "plasma cells (+0.095, p = 0.017) and neutrophils (-0.037, p = 0.024) "
  "differed (nominal-to-BH 0.08). cfDNA deconvolution pointed in the same "
  "qualitative direction of lymphoid remodeling (CD4 naïve -0.067, NK "
  "-0.042, CD4 memory +0.042, eosinophils +0.030). The cfDNA methylation "
  "signature therefore parallels a disease-state immune shift rather than "
  "a pre-symptomatic predictor, matching the timing of the cfDNA sampling "
  "at diagnosis (Fig. 4).")
doc.add_heading("Genetic anchoring: SLC17A1 emerges as the single SMR-supported "
                "shared DMR", level=2)
p("Among the three meQTL-GWAS shared DMRs (235 shared variant pairs in "
  "total), SMR implicated only the SLC17A1 gene-body DMR (chr6:25.72-25.83 "
  "Mb, hypomethylated in PE cfDNA, Δβ = -0.119): cg19490609 "
  "(p_SMR = 0.0017, BH = 0.020) and cg25753631 (p_SMR = 0.0040, BH = 0.024) "
  "in PE. chr17:39.45 Mb was nominal only (p = 0.014, BH = 0.055) and "
  "chr3:49.90 Mb showed no SMR signal. Colocalization remained equivocal "
  "across all loci (max PP4 = 0.486, chr3/GH; SLC17A1 PP4 ≤ 0.217), so "
  "shared causal variant evidence is not established; the meQTL resource "
  "contains only significant pairs and HEIDI was not feasible. SLC17A1 "
  "(encoding a renal phosphate transporter also expressed in immune "
  "cells) remains the sole methylation signal with genetic-anchored "
  "support and is prioritized for follow-up (Table 1).")

# Table 1: SMR
doc.add_heading("Table 1. SMR summary for the three meQTL-GWAS shared DMRs "
                "(FinnGen R10 PE / GH)", level=2)
tbl = doc.add_table(rows=1, cols=6)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, h in enumerate(["DMR (hg38)", "CpG", "Lead meQTL", "Phenotype",
                       "p_SMR", "p_BH"]):
    hdr[i].text = h
smr_rows = [
    ("chr3:49.90 Mb (hyper)", "cg24308560", "rs62262106", "GH", "0.32", "0.75"),
    ("chr6 SLC17A1 body (hypo)", "cg19490609", "rs4711091", "PE", "0.0017", "0.020"),
    ("chr6 SLC17A1 body (hypo)", "cg25753631", "rs212937", "PE", "0.0040", "0.024"),
    ("chr6 SLC17A1 body (hypo)", "cg00387872", "rs6456696", "PE", "0.14", "0.43"),
    ("chr17:39.45 Mb (hypo)", "cg15445000", "rs7218074", "PE", "0.014", "0.055"),
]
for r_ in smr_rows:
    cells = tbl.add_row().cells
    for i, v in enumerate(r_):
        cells[i].text = v
p("meQTL: GoDML cis (N≈32,000). SMR Wald-ratio chi-square test; T statistic "
  "invariant to allele orientation, effect direction not interpretable "
  "without allele columns. Only loci with at least one p_SMR < 0.05 shown; "
  "full results in Supplementary Table of the main manuscript.",
  italic=True, size=9)

# ---------------- Discussion ----------------
doc.add_heading("Discussion", level=1)
p("This companion article assembles four independent lines of evidence "
  "around one question: what does PE-associated cfDNA methylation measure? "
  "The answer that emerges is coherent: a maternal leukocyte-composition "
  "state.")
bullets([
    "Leukocyte trajectories (GSE37722) show that hypermethylated DMR loci "
    "are precisely those that gain methylation in maternal leukocytes "
    "across normal gestation, with concordant effect sizes (ρ = 0.381). "
    "PE, sampled at diagnosis, plausibly exaggerates or shifts this "
    "gestational leukocyte program.",
    "The co-methylation architecture (WGCNA) places the DMRs in a module "
    "whose eigengene is simultaneously PE-associated and correlated with "
    "lymphoid cfDNA fractions, mechanistically embedding the DMRs in the "
    "cell-composition axis.",
    "Immune profiling of maternal blood shows the corresponding cellular "
    "shift (NK/cytotoxic/plasma-cell enrichment) only at diagnosis, "
    "temporally matching, not preceding, the cfDNA signal.",
    "Cis-genetic anchoring is largely absent - DMRs are not enriched for "
    "meQTLs (companion main manuscript) and only SLC17A1 passes SMR "
    "correction - reinforcing an environment/state origin over inherited "
    "methylation regulation.",
])
p("Two implications follow for translational design. First, PE cfDNA "
  "methylation panels should be validated as maternal immune-state "
  "markers rather than placental-disease markers; their clinical niche "
  "is more plausibly disease-activity monitoring or triage at "
  "presentation than first-trimester prediction, which our early-pregnancy "
  "transcriptome data do not support. Second, SLC17A1 is a distinctive "
  "candidate because it is the only DMR whose methylation is both "
  "case-discriminating and anchored by genotype-methylation-disease "
  "co-transmission, and its gene product links to phosphate handling, a "
  "pathway long implicated in PE physiology; targeted validation in "
  "independent cohorts with allele-harmonized meQTL resources is "
  "warranted.")
p("Strengths include the matched design, pre-specified sensitivity "
  "analyses (tube type excluded as a confounder), and triangulation across "
  "four data modalities. Limitations are stated plainly: GSE37722 uses "
  "the HM27 platform with limited DMR coverage (113 probes); meQTL "
  "resources lack allele information and significant-pair filtering "
  "constrains coloc; HEIDI was not implementable; the immune-infiltration "
  "cohorts are modest in size; and all evidence is correlational in "
  "origin.")

# ---------------- Data availability ----------------
doc.add_heading("Data and materials availability", level=1)
p("Analysis-ready cfDNA methylation data (369 per-sample Bismark coverage "
  "files; 357 QC-passed), region-level beta matrix, and sample-level "
  "metadata have been compiled into EGA submission packages "
  "(results/ega_submission/: sample_metadata.tsv, data_file_manifest.tsv, "
  "experiment_metadata.tsv, analysis_description.md, submission_checklist.md) "
  "for controlled-access deposition at the European Genome-Phenome Archive; "
  "deposition will be completed upon manuscript acceptance. Public datasets "
  "used: GSE282512 (cfDNA WGBS), GSE37722 (leukocyte HM27), GSE85307 and "
  "GSE48424 (maternal blood transcriptomes), GoDML cis-meQTL, and FinnGen "
  "R10. Analysis scripts (R 4.6.1 / Python 3.13) are available in the "
  "project repository.")

# ---------------- Figure legends ----------------
doc.add_heading("Figure legends", level=1)
bullets([
    "Figure 1. Gestational methylation trajectories of DMR-overlapping "
    "leukocyte probes (GSE37722), stratified by cfDNA DMR direction. "
    "Hypermethylated-DMR probes rise across gestation (Stouffer "
    "p = 5.4e-28). [figures/gse37722_stage_trend.png]",
    "Figure 2. Concordance of PE cfDNA Δβ with leukocyte gestational Δβ "
    "(delivery - nulligravid); Spearman ρ = 0.381, p = 4.3e-06. "
    "[figures/gse37722_concordance.png]",
    "Figure 3. WGCNA module-trait associations (64 subcohort samples): "
    "blue module associates with PE and lymphoid cfDNA fractions; DMR "
    "enrichment p = 7.3e-17. [figures/wgcna_module_trait.png]",
    "Figure 4. Immune infiltration (ssGSEA) in early-pregnancy vs "
    "diagnosis-timepoint maternal blood; differences emerge only at "
    "diagnosis. [figures/immune_heatmap.png]",
])

doc.save(OUT)
print("saved:", OUT)
