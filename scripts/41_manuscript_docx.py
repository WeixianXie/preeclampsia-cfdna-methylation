# -*- coding: utf-8 -*-
# 41_manuscript_docx.py — C: 投稿文稿汇总 (Methods/Results/讨论+limitations) 生成 docx
import csv, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

def read_perf():
    """读取 B 分类器性能表 (model -> dict)"""
    rows = []
    p = os.path.join(ROOT, "results", "GSE282512_classifier_performance.csv")
    if os.path.exists(p):
        with open(p, encoding="utf-8") as f:
            rows = list(csv.DictReader(f))
    return {r["model"]: r for r in rows}

PERF = read_perf()
def g(model, key, fmt="{:.3f}"):
    r = PERF.get(model)
    if not r or r.get(key) in ("", None, "NA", "NA_real_"):
        return "NA"
    try:
        return fmt.format(float(r[key]))
    except ValueError:
        return r.get(key, "NA")

doc = Document()
# 全局样式
st = doc.styles["Normal"]
st.font.name = "Times New Roman"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
for i, sz, col in [(1, 16, "000000"), (2, 13, "000000"), (3, 11.5, "1F4E79")]:
    h = doc.styles[f"Heading {i}"]
    h.font.name = "Times New Roman"; h.font.size = Pt(sz); h.font.color.rgb = RGBColor.from_string(col)
    h.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

def para(text, bold=False, italic=False, size=None, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def rich(parts, space_after=6):
    """parts: list of (text, bold) 或纯字符串(默认非粗体)"""
    p = doc.add_paragraph()
    for it in parts:
        t, b = it if isinstance(it, tuple) else (it, False)
        r = p.add_run(t); r.bold = b
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet").paragraph_format.space_after = Pt(3)

def table(header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]; cell.text = h
        for r in cell.paragraphs[0].runs: r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ============ 标题与摘要 ============
para("妊娠期高血压疾病血浆游离 DNA 甲基化组学特征：差异区域图谱、母体来源机制证据与诊断分类器性能",
     bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("（投稿文稿汇总草稿：Methods / Results / Discussion & Limitations）",
     italic=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

doc.add_heading("摘要", level=1)
para("目的 基于妊娠期高血压疾病（HDP，重点为子痫前期 PE）患者血浆游离 DNA（cfDNA）全基因组亚硫酸氢盐测序（WGBS）数据，刻画差异甲基化区域（DMR）图谱，系统检验其分子来源机制，并评估其作为无创诊断生物标志物的分类性能。"
     "方法 使用 GSE282512 队列（279 例患者、369 份血浆 cfDNA WGBS 样本，QC 通过 357 份；孕周匹配子队列 PE 32 例 vs 对照 32 例），以覆盖度加权的区域级甲基化为特征、limma 调整采血管与批次效应检出 DMR；通过 GoDMC cis-meQTL 与 FinnGen PE GWAS 进行遗传层整合；通过跨队列表达验证、调控元件富集、胎盘组织甲基化对照与 JASPAR 转录因子 motif 富集进行机制定位；以 EpiDISH 白细胞判别标记集评分与去卷积直接检验细胞构成假设；以留一交叉验证（LOOCV）弹性网逻辑回归评估 DMR 分类器性能并辅以标签置换检验。"
     f"结果 共检出 166 个候选 DMR（高甲基化 105 / 低甲基化 61），信号在 EDTA 单管子集（方向一致率 100%，63 个区域 FDR 复现）与早发型 PE（EOPE）子集中富集；GO 通路富集无 BH 显著条目，名义信号集中于高甲基化子集的有丝分裂纺锤体相关过程。分类器在子队列内 LOOCV AUC 达 {g('elastic_net_166DMR','auc')}（95% CI {g('elastic_net_166DMR','auc_lo')}–{g('elastic_net_166DMR','auc_hi')}），全部 20 次标签置换的 AUC（最高 0.791）均低于观测值；校准斜率 0.923、Hosmer-Lemeshow p=0.056，决策曲线在 0.10–0.50 阈值区间净获益为正；top-20 精简 panel AUC {g('top20_foldselect','auc')}，LOPE 亚组 AUC 最高（{g('elastic_net_LOPE_vs_CT','auc')}）。"
     "机制整合呈四重一致证据链：顺式表达验证（8 队列 meta）为阴性；DMR 富集组织判别性增强子（OR=1.47）而缺失于启动子（OR=0.48）；cfDNA 甲基化变化与胎盘组织变化呈负相关（ρ=-0.522）；PE cfDNA 谱显著偏离胎盘谱（ρ=-0.669）。JASPAR 富集显示髓系炎症 AP-1 程序（FOS OR=15.3）激活而缺氧/发育程序（ARNT::HIF1A）缺失。直接白细胞亚群去卷积为阴性（受稀疏覆盖功效限制）。"
     "结论 PE 相关 cfDNA DMR 是母体血液学/炎症改变的表观遗传指纹而非胎盘甲基化病变的直接释放；其组合可作为候选无创母体生物标志物，但需前瞻队列与独立队列验证。")

# ============ 引言 ============
doc.add_heading("1 引言（要点）", level=1)
bullets([
    "PE 的分子监测目前依赖胎盘来源标志物（如 sFlt-1/PlGF、胎盘 cfDNA 片段组学），母体血管内皮与血液学改变缺乏无创表观遗传标志物。",
    "血浆 cfDNA 甲基化组可同时携带组织来源与细胞构成信息，但 WGBS 稀疏覆盖下单 CpG 分析功效有限，区域级聚合的生物学解释存在争议。",
    "本研究在matched子队列中构建 DMR 图谱后，围绕三个问题层层推进：① 信号是否稳健（采血管/批次/亚型敏感性）；② 信号来源为何（胎盘 vs 母体；遗传调控 vs 构成变化）；③ 信号是否具备诊断转化价值（LOOCV 分类器性能与精简 panel）。",
])

# ============ 方法 ============
doc.add_heading("2 方法", level=1)

doc.add_heading("2.1 队列与数据", level=2)
para("数据集：GSE282512（血浆 cfDNA WGBS，hg38，cov 格式：chr/start/end/meth%/countM/countU）。279 例患者共 369 份样本，质控（覆盖深度、CpG 测得数、平均 β 分布）通过 357 份。分析子队列为孕周 ±3 周匹配的 PE 32 例 vs 对照 32 例（中位孕周差 0 周），其中 PE 含 EOPE 21 例、LOPE 10 例、发病分型缺失 1 例。采血管类型（PAXgene/EDTA）在组间不均衡（PE 12/20 vs 对照 7/25），已通过敏感性分析排除其混杂（见 4.1）。")

doc.add_heading("2.2 甲基化预处理与区域级 DMR 检出", level=2)
para("逐样本汇总 cov 的 countM/countU 至注释区域（CPG 岛/shore/shelf/promoter/gene body 等，Zhou HM450 hg38 区域注释并映射至 WGBS 覆盖空间），构建覆盖度加权区域级 β 矩阵。以 limma 拟合 group + tube_type + batch 设计矩阵检出 DMR（Benjamini-Hochberg FDR 控制）。位点级以单 CpG 均值 β 做验证性检验并计算与区域方向的一致率。敏感性分析：① EDTA-only 子集重跑全部检验；② 采血管 × 疾病交互检验；③ EOPE 子集独立分析。候选 DMR 按“位点验证 + 敏感性复现 + EOPE 富集”综合分级（v2 分级：high_confidence / core_tube_robust / tube_robust / sensitivity_only / deprioritized）。")

doc.add_heading("2.3 遗传层整合（meQTL 与 GWAS）", level=2)
para("GoDML cis-meQTL（坐标 GRCh37，经 Ensembl REST map/cgd 回映至 GRCh38 后联表）检验 DMR 是否富集甲基化数量性状位点调控区域；FinnGen PE GWAS 全显著位点检验 lead SNP 邻近关系与区域亚阈值信号（±250 kb，S Lift over/坐标回映后）。三角联表：DMR × meQTL × GWAS 共享变异。")

doc.add_heading("2.4 机制层：表达验证、调控元件与胎盘对照", level=2)
para("① 顺式表达：GSE98224 胎盘转录组主验证（对应基因 DE 检验）+ 8 队列跨组织 meta（胎盘 6 队列 + 母血 2 队列）；② 调控元件：DMR 与 Ensembl regulatory build（增强子/启动子/CTCF/开放染色质） Fisher 富集，以全区域为背景；③ 胎盘组织复现：cfDNA Δβ 与独立胎盘组织甲基化 Δβ 的 Spearman 相关；④ 来源检验：cfDNA Δβ 与（胎盘 β − 对照 cfDNA β）偏差评分相关。")

doc.add_heading("2.5 转录因子 motif 富集（JASPAR）", level=2)
para("JASPAR2024 CORE 人 PWM 集（1,661 个），Biostrings 双链扫描候选 DMR 序列与 GC 含量 + 长度匹配的背景区域（2,458 个），min.score=80%，Fisher 检验 + BH 校正。")

doc.add_heading("2.6 白细胞亚群去卷积与判别标记集评分", level=2)
para("EpiDISH centDHSbloodDMC.m 333 个跨平台白细胞亚型判别标记（7 亚型：B/NK/CD4T/CD8T/Mono/Neutro/Eosino）与 cent12CT450k.m（12 亚型）。鉴于 cov 位置 = Zhou manifest CpG_beg + 1（经位置集合交集与参考 β 双重校准），自 cov 流式抽取每标记 ±250 bp 窗口内 countM/countU（awk 归并扫描，O(n)）。分析：① 判别标记集评分——每亚型取参考 β<0.2 且其余亚型 max>0.6 的标记，样本内均值 β 越低提示该亚型来源占比越高（免矩阵求逆，规避共线病态）；② RPC/C/CP 去卷积（覆盖 ≥50% 标记，行中位数插补）；③ Wilcoxon + lm（pe+tube+batch）组间检验；④ 评分与 DMR 甲基化的 Spearman 联动。")

doc.add_heading("2.7 DMR 基因功能富集（GO）", level=2)
para("候选 DMR 映射基因：直接符号注释基因 + 无符号区域（regbuild 增强子等）按区间中点到基因体最小距离取最近基因（166 区域 → 155 基因，direct 130 / nearest 36）。以全部检验区域所注释的 35,143 基因为背景宇宙，GO 三大本体（BP/CC/MF，限定背景内 10–500 基因的条目，含祖先传播，过滤 ND 证据）做超几何检验 + BH 校正，hyper/hypo/all 三套分别富集。KEGG 通路富集因 KEGG REST 服务在本分析环境不可达而未执行。")

doc.add_heading("2.8 分类器构建、性能与深度评估", level=2)
para("特征：166 个候选 DMR 的区域级 β（NA 率 1.6%）。模型：弹性网逻辑回归（α=0.5）。评估：① 留一交叉验证（LOOCV），每次 fold 内完成 NA 行中位数插补、5 折内层 λ 选择（lambda.min），无信息泄漏；② 20 次标签置换的 LOOCV AUC 零分布；③ top-k（k=10/20）精简 panel——每 fold 内按 |t| 统计量重选特征（岭回归拟合）；④ EOPE vs 对照、LOPE vs 对照子集内独立 LOOCV；⑤ 仅含采血管与批次的协变量 sanity 模型；⑥ 对照模型——单变量方向加权评分 + 单变量逻辑回归、SVM（RBF 核、概率输出）、XGBoost（fold 内 xgb.cv 定轮数），均同协议 LOOCV。深度评估：rms::val.prob 分组校准（校准截距/斜率 + 不可靠性 U 检验）、Hosmer-Lemeshow 检验（10 分组）、dcurves 决策曲线分析（DCA，净获益）、XGBoost SHAP 特征重要性（predcontrib 原生 TreeSHAP）、以及 6 个 LOOCV 100% 选择频率区域的 rms 列线图（lrm）。指标：AUC（DeLong 95% CI）、Youden 阈值下的灵敏度/特异度/准确率、Brier 评分。pROC、glmnet、xgboost、e1071、rms、dcurves、ResourceSelection 实现。")

doc.add_heading("2.9 软件与统计", level=2)
para("R 4.6.1（data.table、limma、EpiDISH、glmnet、pROC、ggplot2；UTF-8 编码约定）；Ensembl REST；多重检验一律 BH 校正；显著性阈值 FDR<0.05（富集与去卷积层同）。全部分析脚本与结果表随文稿归档于项目 results/ 与 scripts/ 目录。")

# ============ 结果 ============
doc.add_heading("3 结果", level=1)

doc.add_heading("3.1 DMR 图谱与稳健性", level=2)
rich([("共检出 166 个候选 DMR（高甲基化 105，低甲基化 61）。", True),
      ("最强信号为 chr12 LOC124902940 启动子区（Δβ=0.145，hyper），同时通过位点验证、EDTA 敏感性与 EOPE 三重检验。", False)])
para("位点级验证：无一单 CpG 过 FDR，与区域方向一致率中位 44%——提示 cfDNA 覆盖度构成效应（片段来源混合）稀释了单点信号，区域级聚合更稳健。分级：site_confirmed 6 / direction_consistent 39 / unconfirmed 121；v2 综合分级 high_confidence 6 / core_tube_robust 7 / tube_robust 72 / sensitivity_only 30 / deprioritized 51。")
para("敏感性：EDTA-only 子集方向一致率 100%（0 翻转）、63 个候选 FDR 复现、管型×疾病交互检验全部不显著——采血管混杂被排除。EOPE 子集方向一致 100%、27 个 FDR 复现，信号在早发型富集。")

doc.add_heading("3.2 遗传层整合：DMR 独立于经典易感位点", level=2)
para("候选 DMR 不富集 GoDMC cis-meQTL（68.7% vs 背景 75.7%，OR=0.70，Fisher p=0.985；meQTL 密度更低，Wilcoxon p=0.027）；不富集 FinnGen PE 全显著位点邻近区域（0/166，p=1），但 50/166 存在区域亚阈值信号。三角联表（坐标回映修正后）检出 235 对 meQTL-GWAS 共享变异，涉及 3 个 DMR：chr6 SLC17A1 基因体（hypo，rs116510165 PE p=4.3×10⁻⁵）、chr3:49.9 Mb、chr17:39.45 Mb——为甲基化-基因型-疾病三重交汇候选。")

doc.add_heading("3.3 机制层：母体来源的四重一致证据", level=2)
bullets([
    "顺式表达验证阴性：GSE98224 主验证 0/63 基因 FDR 复现；8 队列 meta（胎盘 6 + 母血 2）0 显著、胎盘方向一致 0/64；SLC17A1 八队列近零——排除顺式基因调控解释。",
    "调控元件：DMR 富集增强子（OR=1.47，BH=0.037）而缺失于启动子（OR=0.48，BH=3.4×10⁻⁴）——落在组织判别性调控元件而非管家启动子。",
    "胎盘组织复现阴性且反向：cfDNA Δβ vs 胎盘组织 Δβ Spearman ρ=-0.522（p=0.0026）。",
    "来源检验：PE cfDNA 谱显著偏离胎盘谱（vs 胎盘−对照 cfDNA 偏差评分，ρ=-0.669，p=3.9×10⁻⁵）。",
])
para("JASPAR motif：FOS/AP-1 显著富集（OR=15.3，BH=0.022；JUN/JUND/FLI1/ETV3/IRF1/YY1 名义富集）；ARNT::HIF1A（OR=0.50，BH=0.018）、CGGBP1（BH=0.014）显著缺失；hypo 子集 LMX1B 缺失（BH=0.004）；滋养层因子 TEAD4 无富集（OR=0.78）。总体指向髓系炎症程序激活、缺氧/发育程序缺失的母体炎症表观指纹。")

doc.add_heading("3.4 白细胞亚群构成的直接检验（阴性，功效受限）", level=2)
para("333 个判别标记在 64 份子队列样本中并集可测 281 个（84%），样本级中位可测 72/333（21.6%），平均可测率 26.5%（WGBS 稀疏覆盖所限）。7 亚型判别标记集评分组间差异均不显著（最显著 Neutro，p=0.196，BH=0.94）；12 亚型仅 Mono 达最低标记数且不显著；RPC/CP 去卷积全不显著（RPC 下 Eosino/CD4T 分数为 0，共线性不稳）。评分与 DMR 甲基化联动仅 CD8T×低甲基化 DMR 名义相关（ρ=0.32，p=0.010，BH=0.14）。")
rich([("解释：区域注释级（中位宽 18 kb）去卷积曾提示 NK/CD4Tnv 降低，但在标记级未复现，判定为区域稀释假象。", True),
      ("直接证据不支持白细胞亚群构成显著偏移，但 26.5% 可测率与 n=32/32 下的统计功效不足以排除中度差异；“构成指纹”仍为机制推断，需粒细胞亚群靶向标记验证（如流式分选参考谱或靶向测序 panel）。")])

doc.add_heading("3.5 DMR 基因功能富集（GO）", level=2)
para("166 个候选 DMR 映射至 155 个基因（直接注释 130，最近基因 36），以 35,143 个检验宇宙基因为背景做 GO 三本体富集（BP 5,707 / CC 739 / MF 1,222 个可检条目）。无任何条目通过 BH FDR<0.05（all/hyper/hypo 三套一致）——与 2.6 节白细胞亚群去卷积阴性相呼应：DMR 信号不集中在特定经典免疫通路。名义显著（p<0.01）条目集中于高甲基化子集的有丝分裂纺锤体与微管组织（mitotic spindle、establishment of mitotic spindle orientation、regulation of microtubule nucleation）及 rRNA 加工等增殖相关过程，方向上与髓系细胞活化/增殖状态一致，但因未过多重校正仅作探索性提示（全表见补充表 S7/S19）。")

doc.add_heading("3.6 分类器性能、校准与临床效用", level=2)
rich([(f"166 DMR 弹性网分类器 LOOCV AUC = {g('elastic_net_166DMR','auc')}", True),
      (f"（95% CI {g('elastic_net_166DMR','auc_lo')}–{g('elastic_net_166DMR','auc_hi')}），"
       f"灵敏度 {g('elastic_net_166DMR','sens')}、特异度 {g('elastic_net_166DMR','spec')}、准确率 {g('elastic_net_166DMR','acc')}。")])
para(f"20 次标签置换零分布 AUC 均值 0.632、最大 0.791，全部低于观测值（置换 p 触及 20 次检验的理论下界 1/21≈0.048）——分类性能非标签噪声或过拟合所致。仅含采血管与批次的协变量模型 AUC=" + g("tube_batch_only","auc") + "，排除混杂驱动。")
rows = [
    ["全 166 DMR 弹性网", "64 (32/32)", g('elastic_net_166DMR','auc'), f"{g('elastic_net_166DMR','auc_lo')}–{g('elastic_net_166DMR','auc_hi')}", g('elastic_net_166DMR','sens'), g('elastic_net_166DMR','spec')],
    ["top-10 panel（fold 内选择）", "64 (32/32)", g('top10_foldselect','auc'), f"{g('top10_foldselect','auc_lo')}–{g('top10_foldselect','auc_hi')}", g('top10_foldselect','sens'), g('top10_foldselect','spec')],
    ["top-20 panel（fold 内选择）", "64 (32/32)", g('top20_foldselect','auc'), f"{g('top20_foldselect','auc_lo')}–{g('top20_foldselect','auc_hi')}", g('top20_foldselect','sens'), g('top20_foldselect','spec')],
    ["EOPE vs 对照（子集 LOOCV）", "53 (21/32)", g('elastic_net_EOPE_vs_CT','auc'), f"{g('elastic_net_EOPE_vs_CT','auc_lo')}–{g('elastic_net_EOPE_vs_CT','auc_hi')}", "—", "—"],
    ["LOPE vs 对照（子集 LOOCV）", "42 (10/32)", g('elastic_net_LOPE_vs_CT','auc'), f"{g('elastic_net_LOPE_vs_CT','auc_lo')}–{g('elastic_net_LOPE_vs_CT','auc_hi')}", "—", "—"],
    ["采血管+批次（sanity）", "64 (32/32)", g('tube_batch_only','auc'), "—", "—", "—"],
]
table(["模型", "n（PE/CT）", "AUC", "95% CI", "灵敏度", "特异度"], rows)
para(f"精简 panel 保持高性能（top-20 AUC {g('top20_foldselect','auc')}），提示可通过 fold 内一致性选择压缩标志物数量。LOPE 亚组 AUC 最高（{g('elastic_net_LOPE_vs_CT','auc')}）而 EOPE 较低（{g('elastic_net_EOPE_vs_CT','auc')}）——与 DMR 信号在 EOPE 富集的发现互补，可能与 LOPE 样本量小（n=10）导致区间不稳定有关，需独立队列复核。")
para("模型深度评估：校准良好——LOOCV 预测的校准斜率 0.923（理想=1）、截距 -0.243、不可靠性检验 U:p=0.80，Brier 评分 0.095；Hosmer-Lemeshow 检验（10 分组）χ²=15.17、p=0.056（未拒绝校准良好）。决策曲线分析显示分类器在 0.10–0.50 阈值区间净获益 0.46→0.36，优于全治策略（患病率 0.5）。对照模型（同协议 LOOCV）：单变量评分逻辑回归 AUC=0.895、SVM（RBF）AUC=0.921、XGBoost AUC=0.802——弹性网（0.932）在高维稀疏特征下最优。SHAP 重要性（全数据 XGBoost）top 区域为 73424/63782/63444/31006/612；6 个 LOOCV 100% 选择频率区域构建的列线图表现观 C-index=0.972（样本内，用于展示评分形式而非外推性能）。")

# ============ 讨论 ============
doc.add_heading("4 讨论与局限性", level=1)
para("综合证据链：cfDNA DMR 信号在区域级稳健（覆盖度加权、管型/批次调整、EDTA 敏感性全过），但独立于经典遗传易感位点与 cis-meQTL 调控、无顺式表达后果、与胎盘组织甲基化变化反向、偏离胎盘谱，并携带髓系炎症 AP-1 程序的 motif 签名——最佳解释为母体血细胞（白细胞亚群比例或活化状态）变化的 cfDNA 表观指纹，而非胎盘滋养层甲基化病变的直接释放。在此框架下，DMR 组合的分类性能（LOOCV AUC≈0.93）反映的是 PE 母体炎症/血液学改变的可检测性，与胎盘标志物通路形成互补。")
doc.add_heading("4.1 局限性", level=2)
bullets([
    "单队列、回顾性设计：全部结果源自 GSE282512 单中心队列，分类器无外部独立验证；LOOCV 性能存在留一乐观性（置换检验已部分缓解）。",
    "样本量：子队列 32/32（EOPE 21、LOPE 10），亚组分析区间宽；1 例 PE 发病分型缺失。",
    "覆盖稀疏：WGBS 平均 ~17% CpG、标记级 26.5% 可测率，单 CpG 验证与白细胞去卷积功效受限；构成假设未被直接证实，也不能排除粒细胞活化（而非亚群比例）等更细粒度机制。",
    "采血管不均衡已通过敏感性分析排除混杂，但 PAXgene 对 cfDNA 谱的微小影响不能完全排除（残余混杂风险低）。",
    "匹配对照为孕周匹配而非同龄健康对照孕妇以外的其他产科并发症对照，缺乏疾病特异性对照（如 FGR、慢性高血压）。",
    "三角联表的 3 个共享 DMR（含 SLC17A1）为亚阈值遗传信号，因果方向未定；表达层八队列 meta 为跨平台整合，存在队列间异质性。",
    "分类器阈值（Youden 点）在 OOF 预测上选取，报告的灵敏度/特异度存在轻度乐观；临床转化需前瞻、多中心、预注册验证。列线图 C-index（0.972）为样本内表现，仅演示评分形式。",
    "功能富集规模有限：155 个映射基因在严格 BH 校正下无显著通路，最近基因注释（36 个）为近似映射；KEGG 富集因数据源网络不可达未执行——通路层面结论以探索性表述为主。",
])
doc.add_heading("4.2 后续工作", level=2)
bullets([
    "独立队列外部验证分类器（优先靶向 panel：top-20 DMR 区域富集扩增）。",
    "粒细胞亚群/活化状态的定向表观验证（流式分选参考谱 or 靶向甲基化测序）。",
    "SLC17A1 等三重交汇位点的功能专题（基因型-甲基化-表达的孟德尔随机化式分析）。",
    "与胎盘标志物（sFlt-1/PlGF、胎盘 cfDNA 片段组学）的联合模型评估增量诊断价值。",
])

# ============ 附：图表与文件清单 ============
doc.add_heading("附录 A 图表清单", level=1)
table(["编号", "内容", "文件"], [
    ["图 1", "分类器 LOOCV ROC 曲线", "figures/roc_loocv.png"],
    ["图 2", "校准曲线（LOOCV 预测）", "figures/calibration.png"],
    ["图 3", "决策曲线分析（DCA）", "figures/dca.png"],
    ["图 4", "SHAP 特征重要性 top-20", "figures/shap_importance.png"],
    ["图 5", "6 稳定区域列线图", "figures/nomogram.png"],
    ["图 6", "JASPAR TF 富集 top", "figures/jaspar_top_tf.png"],
    ["图 7", "GO 富集气泡图（名义 top）", "figures/enrichment_bubble.png"],
    ["图 8", "去卷积/细胞评分箱线图", "figures/cellscore_box.png"],
    ["表 S1", "DMR v2 分级全表", "results/GSE282512_dmr_final_v2.csv"],
    ["表 S2", "分类器性能", "results/GSE282512_classifier_performance.csv"],
    ["表 S3", "panel 选择频率", "results/GSE282512_classifier_panel.csv"],
    ["表 S4", "置换零分布", "results/GSE282512_classifier_perm_null.csv"],
    ["表 S5", "白细胞评分 7 亚型", "results/GSE282512_cellscore_7type.csv"],
    ["表 S6", "JASPAR 富集全表", "results/GSE282512_jaspar_motif_enrichment.csv"],
    ["表 S7", "GO 富集全表", "results/GSE282512_dmr_enrichment_go.csv"],
    ["表 S8", "调控元件富集", "results/GSE282512_regbuild_enrichment.csv"],
    ["表 S9", "LOOCV 样本级预测值", "results/GSE282512_classifier_oof_predictions.csv"],
    ["表 S10", "对照模型比较", "results/GSE282512_classifier_model_comparison.csv"],
    ["表 S11", "DCA 数据", "results/GSE282512_classifier_dca.csv"],
    ["表 S12", "SHAP 重要性", "results/GSE282512_classifier_shap_importance.csv"],
])
para("补充材料（S1–S8 全部汇总文本、完整补充图表清单、代码与数据可用性声明）见 results/GSE282512_supplementary_materials.docx。")

out = os.path.join(ROOT, "results", "GSE282512_manuscript_draft.docx")
doc.save(out)
print("saved:", out)
