# -*- coding: utf-8 -*-
# 44_supplementary_docx.py — Phase 12: 投稿补充材料 + 代码可用性声明生成
# 动态读取 results/ 各 summary 文件，嵌入 S1-S9 分节
import os, csv, datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = r"E:\妊娠期高血压甲基化研究方案\hdp-methylation-project"
RES = os.path.join(ROOT, "results")
FIG = os.path.join(ROOT, "figures")

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(
    st.element.rPr.rFonts.__class__("w:eastAsia") if False else "w:eastAsia", "宋体") if False else None

def para(t, bold=False, size=10.5, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def mono(t):
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.name = "Consolas"; r.font.size = Pt(8.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def embed(path, maxlines=60):
    fp = os.path.join(RES, path)
    if not os.path.exists(fp):
        para(f"[缺失: {path}]", bold=True); return
    with open(fp, encoding="utf-8", errors="replace") as f:
        lines = [l.rstrip() for l in f][:maxlines]
    for l in lines:
        mono(l if l.strip() else "")

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)

def table(hdr, rows):
    tb = doc.add_table(rows=1, cols=len(hdr)); tb.style = "Light Grid Accent 1"
    for i, hcell in enumerate(hdr):
        tb.rows[0].cells[i].text = str(hcell)
    for row in rows:
        cells = tb.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    para("")

# ---------------- 封面 ----------------
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Supplementary Materials"); r.bold = True; r.font.size = Pt(20)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Differentially methylated regions in maternal plasma cell-free DNA\n"
              "as epigenetic fingerprints of preeclampsia (GSE282512)"); r.font.size = Pt(12)
para(f"生成日期: {datetime.date.today().isoformat()}  |  分析目录: hdp-methylation-project", space_after=14)

# ---------------- S0 补充图清单 ----------------
h1("S0  补充图表清单")
figs = sorted(f for f in os.listdir(FIG) if f.endswith(".png"))
table(["编号", "文件", "说明"], [[f" Supplementary Figure {i+1}", f, ""]
      for i, f in enumerate(figs)])
tbls = [
    ("GSE282512_dmr_final_v2.csv", "候选 DMR v2 综合分级全表 (166 区域)"),
    ("GSE282512_dmr_site_level.csv", "位点级验证结果"),
    ("GSE282512_dmr_sensitivity.csv", "EDTA/PAXgene/交互敏感性"),
    ("GSE282512_phase2_dmr_meqtl.csv", "DMR × meQTL 富集"),
    ("GSE282512_phase3_gwas_overlap.csv", "DMR × GWAS 重叠"),
    ("phase3_shared_mqtl_gwas.csv", "三角联表共享变异"),
    ("GSE282512_expr_meta.csv", "8 队列表达 meta"),
    ("GSE282512_regbuild_enrichment.csv", "调控元件富集"),
    ("GSE282512_placenta_tissue.csv", "胎盘组织复现"),
    ("GSE282512_placenta_origin.csv", "胎盘来源检验"),
    ("GSE282512_jaspar_motif_enrichment.csv", "JASPAR TF motif 富集"),
    ("GSE282512_deconv_fractions.csv", "白细胞去卷积（区域级, 保留供复核）"),
    ("GSE282512_cellscore_7type.csv", "判别标记集评分 7 亚型"),
    ("GSE282512_classifier_performance.csv", "分类器性能汇总"),
    ("GSE282512_classifier_oof_predictions.csv", "LOOCV 样本级预测值"),
    ("GSE282512_classifier_model_comparison.csv", "对照模型比较"),
    ("GSE282512_classifier_dca.csv", "DCA 决策曲线数据"),
    ("GSE282512_classifier_shap_importance.csv", "SHAP 特征重要性"),
    ("GSE282512_dmr_enrichment_go.csv", "DMR 基因 GO 富集"),
    ("GSE282512_manuscript_draft.docx", "主文稿"),
]
table(["编号", "文件", "说明"],
      [[f" Supplementary Table {i+1}", f, d] for i, (f, d) in enumerate(tbls)])

# ---------------- S1 数据与 QC ----------------
h1("S1  数据集与质量控制")
embed("GSE282512_qc_summary.csv", maxlines=10)
embed("GSE282512_subcohort_summary.txt")

# ---------------- S2 DMR ----------------
h1("S2  DMR 发现、位点验证与敏感性")
embed("GSE282512_dmr_final_summary.txt")
embed("GSE282512_dmr_sensitivity_summary.txt")
embed("GSE282512_dmr_final_v2_summary.txt")

# ---------------- S3 遗传层 ----------------
h1("S3  遗传层整合 (meQTL / GWAS / 三角联表)")
embed("GSE282512_phase2_summary.txt")
embed("GSE282512_phase3_summary.txt")
embed("phase3_shared_summary.txt")

# ---------------- S4 机制层 ----------------
h1("S4  机制层 (表达 meta / 调控元件 / 胎盘对照 / JASPAR)")
embed("GSE282512_expr_validation_summary.txt")
embed("GSE282512_expr_meta_summary.txt")
embed("GSE282512_regbuild_enrichment.txt")
embed("GSE982512_expr_validation_summary.txt")
embed("GSE282512_jaspar_summary.txt")

# ---------------- S5 细胞构成 ----------------
h1("S5  白细胞亚群去卷积与标记级评分")
embed("GSE282512_deconv_summary.txt")
embed("GSE282512_cellscore_summary.txt")

# ---------------- S6 分类器 ----------------
h1("S6  分类器性能、校准与对照模型")
embed("GSE282512_classifier_summary.txt")
embed("GSE282512_classifier_depth_summary.txt")

# ---------------- S7 功能富集 ----------------
h1("S7  DMR 基因功能富集 (GO)")
embed("GSE282512_dmr_enrichment_summary.txt")

# ---------------- S8 代码与数据可用性 ----------------
h1("S8  Code and Data Availability")
para("Data availability", bold=True)
para("All datasets are publicly available: GSE282512 (plasma cfDNA WGBS), "
     "GSE98224 et al. (transcriptome cohorts), placenta methylation cohorts (see S4), "
     "GoDMC mQTL and FinnGen R10 GWAS summary statistics via respective public portals. "
     "No new data were generated.")
para("Code availability", bold=True)
scripts = sorted(f for f in os.listdir(os.path.join(ROOT, "scripts"))
                 if f[0].isdigit() and (f.endswith(".R") or f.endswith(".sh") or f.endswith(".py")))
para(f"All analysis scripts ({len(scripts)} files, R 4.6.1 / bash / python) are archived in the "
     "project repository and will be deposited on GitHub with a Zenodo DOI upon submission. "
     "Script inventory:")
mono("; ".join(scripts))

para("Archive checklist (upon submission)", bold=True)
for t in [
    "1. Create public GitHub repo (MIT license) containing scripts/ and results/ (summary tables only)",
    "2. Zenodo deposit -> DOI; cite DOI in manuscript Data Availability statement",
    "3. Exclude per-sample raw cov files (>10GB) and _tmp intermediates from the archive",
    "4. Include sessionInfo() output and R package version table in repo README",
    "5. Target journals (per feasibility report): Clinical Epigenetics or J Transl Med (first submission); "
    "Pregnancy Hypertension (fallback)",
]:
    para(t)

out = os.path.join(RES, "GSE282512_supplementary_materials.docx")
doc.save(out)
print("saved:", out)
